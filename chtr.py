#### This File is in charge to read all word .docx files from a subdirectory from your working project
#### directory and process a summary from python nltk processing
import os
import glob
from docx import Document
from docx.oxml import OxmlElement
 
#scrapea articulo de wikipedia
enlace = "AnexoFichasWord"


# Obtener la ruta completa del subdirectorio
ruta_subdirectorio = os.path.join(os.getcwd(), enlace)

# Obtener la lista de archivos en el subdirectorio
archivos = glob.glob(os.path.join(ruta_subdirectorio, '*'))
for docx_file in archivos:
    #Nombre del archivo    
    nombre_archivo = os.path.basename(docx_file)
    print(docx_file)
    # Extraer texto del archivo Word
    #Agregando las tablas del archivo a la sintesis
    tdoc= Document(docx_file)
    for t,tabla in enumerate(tdoc.tables):
        ws=1
        for i, row in enumerate(tabla.rows):
            try:
                len(row.cells)
            except IndexError:
                # Este bloque se ejecutará si se produce una excepción ZeroDivisionError
                print('tabla '+str(t)+'/'+str(len(tdoc.tables))+' doc:'+nombre_archivo)
            else:
                if len(row.cells) <= 0:
                    ws=ws*0
        # Omitir las tablas que tienen columnas combinadas
        if ws==0:
            print('1 tabla oviada')
            continue