#### This File is in charge to read all word .docx files from a subdirectory from your working project
#### directory and process a summary from python nltk processing. It als copy table information
import os
import glob
import docx2txt
import bs4 as bs  
import urllib.request  
import re
import nltk
import bs4
import urllib.request
import requests
import shutil
from bs4 import BeautifulSoup
import urllib.request
from inscriptis import get_text
from googletrans import Translator
from docx import Document
from docx.oxml import OxmlElement
 
#scrapea articulo de wikipedia
enlace = "AnexoFichasWord"


# Obtener la ruta completa del subdirectorio
ruta_subdirectorio = os.path.join(os.getcwd(), enlace)

# Obtener la lista de archivos en el subdirectorio
archivos = glob.glob(os.path.join(ruta_subdirectorio, '*'))
Final_Synth=Document()
Final_Synth.add_heading('Resumen Breve de las acciones de Eficiencia Energética',level=0)
# Iterar sobre los archivos
for docx_file in archivos:
    #Nombre del archivo    
    nombre_archivo = os.path.basename(docx_file)
    Final_Synth.add_heading(nombre_archivo+':',level=1)
    #documento="/AEE01Transformación Del Mercado De Iluminación.docx"
    #html = urllib.request.urlopen(enlace).read().decode('utf-8')

    # Ruta al archivo Word
    #docx_file = enlace+documento
    print(docx_file)
    # Extraer texto del archivo Word
    text = docx2txt.process(docx_file)
    #text = get_text(html)
    article_text = text
    article_text = article_text.replace("[ edit ]", "")
    print ("###################")
    
    from nltk import word_tokenize,sent_tokenize
    # Removing Square Brackets and Extra Spaces
    article_text = re.sub(r'\[[0-9]*\]', ' ', article_text)  
    article_text = re.sub(r'\s+', ' ', article_text)  
    
    formatted_article_text = re.sub('[^a-zA-Z]', ' ', article_text )  
    formatted_article_text = re.sub(r'\s+', ' ', formatted_article_text)  

    nltk.download('stopwords')
    #EN ESTA PARTE HACE LA TOKENIZACION 
    sentence_list = nltk.sent_tokenize(article_text)  
    
    #EN ESTA PARTE ENCUENTRA LA FRECUENCIA DE CADA PALABRA
    stopwords = nltk.corpus.stopwords.words('spanish')
    
    word_frequencies = {}  
    for word in nltk.word_tokenize(formatted_article_text):  
        if word not in stopwords:
            if word not in word_frequencies.keys():
                word_frequencies[word] = 1
            else:
                word_frequencies[word] += 1
    
    #
    maximum_frequncy = max(word_frequencies.values())
    
    for word in word_frequencies.keys():  
        word_frequencies[word] = (word_frequencies[word]/maximum_frequncy)
    
    #CALCULA LAS FRASES QUE MÁS SE REPITEN
    sentence_scores = {}  
    for sent in sentence_list:  
        for word in nltk.word_tokenize(sent.lower()):
            if word in word_frequencies.keys():
                if len(sent.split(' ')) < 30:
                    if sent not in sentence_scores.keys():
                        sentence_scores[sent] = word_frequencies[word]
                    else:
                        sentence_scores[sent] += word_frequencies[word]
    
    #REALIZA EL RESUMEN CON LAS MEJORES FRASES
    import heapq  
    summary_sentences = heapq.nlargest(7, sentence_scores, key=sentence_scores.get)
    
    summary = ' '.join(summary_sentences)  
    print(summary)  
    Final_Synth.add_paragraph(summary)
    #Final_Synth=Final_Synth+'\n'+nombre_archivo+':\n\n'+summary'\n'

    #Agregando las tablas del archivo a la sintesis
    tdoc= Document(docx_file)
    for t,tabla in enumerate(tdoc.tables):
        ws=1
        for i, row in enumerate(tabla.rows):
            if len(row.cells) <= 0:
                ws=ws*0
        # Omitir las tablas que tienen columnas combinadas
        if ws==0:
            print('1 tabla oviada')
            continue
        else:
            # Clonar la tabla y agregarla al documento final
            Final_Synth.add_heading('tabla número:'+str(t+1),level=2)
            nueva_tabla = Final_Synth.add_table(len(tabla.rows), len(tabla.columns),style=tabla.style)
            for i, row in enumerate(tabla.rows):
                for j, cell in enumerate(row.cells):
                    # Copiar el contenido de la tabla original a la nueva tabla
                    nueva_tabla.cell(i, j).text = cell.text



    # Recorrer todas las figuras en el archivo de Word
    #for figura in tdoc.inline_shapes:
        # Agregar la figura al documento final
        #Final_Synth.add_picture(figura)
            
    # for shape in tdoc.inline_shapes:
    #     if shape.has_picture:
    #         # Get the image from the inline shape
    #         image = shape.image

    #         # Save the image to a file
    #         image_path = os.path.join("path/to/save/image.png")
    #         image.save(image_path)
    #         nueva_img = Final_Synth.add_picture(image_path)
    ###
    #traducir:
    ################
    #translator = Translator()
    #translate = translator.translate(summary, src="en", dest="es")
    #print(translate.text)
    ############################
    #print("Haz mi curso gratis de python:  https://www.udemy.com/course/python-desde-0-para-principiantes/")
# Volver a la carpeta base
os.chdir('..')
# Abrir el archivo en modo escritura
#with open('synthesis.txt', 'w', encoding='utf-8') as archivo:
    # Escribir la variable en el archivo
    #archivo.write(Final_Synth)
#Guardar archivo
Final_Synth.save('synthesis.docx')